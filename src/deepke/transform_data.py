import csv
import json
import xlrd2
import docx

def json2csv(fp_json, fp_csv):
  """
    Standard AE RE. 
    .json -> .csv
  """
  with open(fp_json, encoding='utf-8') as f1:
    lines = json.load(f1)
    fieldnames = lines[0].keys()
    with open(fp_csv, 'w', encoding='utf-8') as f2:
      writer = csv.DictWriter(f2, fieldnames=fieldnames)
      writer.writeheader()
      writer.writerows(lines)

def xlsx2csv(fp_xlsx, fp_csv):
  """
    Standard AE RE. 
    .xlsx -> .csv
  """
  book = xlrd2.open_workbook(fp_xlsx)
  sheet = book.sheet_by_index(0)
  row = sheet.row(0)
  row = [str(r).split(':')[1][1:-1] for r in row]
  fieldnames = dict(zip(row, [1] * len(row))).keys()
  row_datas = []

  for idx in range(1, sheet.nrows):
    row_data = sheet.row(idx)
    process_row = []
    for r in row_data:
      temp_r = str(r).split(':')
      if temp_r[0] == 'number':
        process_row.append(temp_r[1][:-2])
      elif temp_r[0] == 'text':
        process_row.append(temp_r[1][1:-1])
    row_data = dict(zip(row,process_row))
    row_datas.append(row_data)

  with open(fp_csv, 'w', encoding='utf-8') as f:
    writer = csv.DictWriter(f, fieldnames=fieldnames)
    writer.writeheader()
    writer.writerows(row_datas)

def json2txt(fp_json, fp_txt):
  """
    Standard NER. 
    .json -> .txt
  """
  with open(fp_json, encoding='utf-8') as f1:
    lines = json.load(f1)
    symbol = []
    for line in lines:
      sentence = line.get('sentence')
      labels = ['O'] * len(sentence)
      entities = line.get('entities')
      for entity in entities:
        word = entity.get('word')
        label = entity.get('label')
        for idx in range(len(sentence) - len(word)):
          if sentence[idx : idx + len(word)] == word:
            labels[idx] = 'B-' + label
            for i in range(idx + 1 , idx + len(word)):
              labels[i] = 'I-' + label
      symbol.append(labels)
    
    with open(fp_txt, 'w', encoding='utf-8') as f2:
      for idx in range(len(lines)):
        for i in range(len(lines[idx].get('sentence'))):
          
          f2.write(lines[idx].get('sentence')[i] + ' ' + symbol[idx][i] + '\n')
        if idx != len(lines) - 1:
          f2.write('\n')

      
def doc2txt(fp_doc, fp_txt):
  """
    Standard NER. 
    .docx -> .txt
  """
  doc = docx.Document(fp_doc)
  temp_sentence = ""
  d = {}
  sentences = []
  for num in range(len(doc.paragraphs)): 
    text = doc.paragraphs[num].text
    if text.find("Sentence") == 0:
      if temp_sentence != "":
        sentences.append(d)
        d = {}
      temp_sentence = text[9:]
      d['sentence'] = text[9:]
      d['entities'] = []
    else:
      label = text.split(':')[0]
      words = text.split(':')[1]
      words = words.split(',')
      for word in words:
        d['entities'].append({"word":word,"label":label})
  if temp_sentence != "":
    sentences.append(d)

  symbol = []
  for line in sentences:
    sentence = line.get('sentence')
    labels = ['O'] * len(sentence)
    entities = line.get('entities')
    for entity in entities:
      word = entity.get('word')
      label = entity.get('label')
      for idx in range(len(sentence) - len(word)):
        if sentence[idx : idx + len(word)] == word:
          labels[idx] = 'B-' + label
          for i in range(idx + 1 , idx + len(word)):
            labels[i] = 'I-' + label
    symbol.append(labels)
    
  with open(fp_txt, 'w', encoding='utf-8') as f2:
    for idx in range(len(sentences)):
      for i in range(len(sentences[idx].get('sentence'))):      
        f2.write(sentences[idx].get('sentence')[i] + ' ' + symbol[idx][i] + '\n')
      if idx != len(sentences) - 1:
        f2.write('\n')
